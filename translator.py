import os
import re
from googletrans import Translator
from docx import Document
import PyPDF2
import pdfplumber

translator = Translator()

def translate_text(text, dest='vi'):
    """Dịch text sang tiếng Việt"""
    if not text or not text.strip():
        return text
    
    try:
        # Chia nhỏ text nếu quá dài (Google Translate có giới hạn)
        if len(text) > 5000:
            parts = [text[i:i+5000] for i in range(0, len(text), 5000)]
            translated_parts = []
            for part in parts:
                result = translator.translate(part, dest=dest)
                translated_parts.append(result.text)
            return ' '.join(translated_parts)
        else:
            result = translator.translate(text, dest=dest)
            return result.text
    except Exception as e:
        print(f"Lỗi dịch: {e}")
        return text

def translate_srt(input_path, output_path):
    """Dịch file SRT"""
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    translated_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Bỏ qua dòng trống
        if not line:
            i += 1
            continue
        
        # Kiểm tra nếu là số thứ tự
        if line.isdigit():
            index = line
            translated_lines.append(index)
            i += 1
            
            # Đọc timestamp
            if i < len(lines):
                timestamp = lines[i].strip()
                translated_lines.append(timestamp)
                i += 1
            
            # Đọc text (có thể nhiều dòng)
            text_lines = []
            while i < len(lines) and lines[i].strip():
                text_lines.append(lines[i].strip())
                i += 1
            
            if text_lines:
                # Gộp text và dịch
                full_text = ' '.join(text_lines)
                translated_text = translate_text(full_text)
                translated_lines.append(translated_text)
            
            # Thêm dòng trống giữa các entry
            translated_lines.append('')
        else:
            i += 1
    
    # Ghi file output
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(translated_lines))
        if translated_lines and translated_lines[-1]:
            f.write('\n')
    
    return output_path

def translate_docx(input_path, output_path):
    """Dịch file Word"""
    doc = Document(input_path)
    
    # Dịch từng paragraph
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated_text = translate_text(paragraph.text)
            paragraph.text = translated_text
    
    # Dịch text trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        translated_text = translate_text(paragraph.text)
                        paragraph.text = translated_text
    
    doc.save(output_path)
    return output_path

def translate_pdf(input_path, output_path):
    """Dịch file PDF - tạo file Word từ PDF đã dịch"""
    all_text = []
    
    # Đọc text từ PDF
    try:
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)
    except:
        # Fallback to PyPDF2
        with open(input_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)
    
    # Dịch tất cả text
    translated_texts = []
    for text in all_text:
        translated = translate_text(text)
        translated_texts.append(translated)
    
    # Tạo file Word từ text đã dịch
    doc = Document()
    for text in translated_texts:
        doc.add_paragraph(text)
        doc.add_page_break()
    
    # Lưu dưới dạng DOCX
    docx_path = output_path.replace('.pdf', '_vi.docx')
    doc.save(docx_path)
    return docx_path

def translate_file(input_path, output_folder):
    """Dịch file dựa trên extension"""
    filename = os.path.basename(input_path)
    name, ext = os.path.splitext(filename)
    ext = ext.lower()
    
    # Tạo tên file output
    if ext == '.srt':
        output_filename = f"{name}_vi.srt"
        output_path = os.path.join(output_folder, output_filename)
        return translate_srt(input_path, output_path)
    
    elif ext in ['.docx', '.doc']:
        output_filename = f"{name}_vi.docx"
        output_path = os.path.join(output_folder, output_filename)
        return translate_docx(input_path, output_path)
    
    elif ext == '.pdf':
        output_filename = f"{name}_vi.docx"  # PDF sẽ được chuyển thành DOCX
        output_path = os.path.join(output_folder, output_filename)
        return translate_pdf(input_path, output_path)
    
    else:
        raise ValueError(f"Định dạng file không được hỗ trợ: {ext}")

